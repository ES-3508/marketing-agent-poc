import streamlit as st
from docx import Document as DocxDocument
import io
import os
import openai

from langchain_openai import ChatOpenAI
from langchain.chains import RetrievalQA
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_core.prompts import PromptTemplate
from langchain.schema import Document


# Define the CSS to inject
def local_css(file_name):
    with open(file_name) as f:
        st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)


def read_word_document(file_buffer):
    doc = DocxDocument(io.BytesIO(file_buffer.getvalue()))
    data = []
    current_question = None
    answer_lines = []

    paragraphs = [p.text for p in doc.paragraphs]

    questions = [
        "What is the Dilemma? What is the opportunity out there? What is the gap that be filled?",
        "How will your brand solve the Dilemma? What does the brand promise the consumer?",
        "What products/service do you offer to solve this?",
        "Why should people believe it? What makes you credible?",
        "What are the trends in the category that inspired you?",
        "Who inspires you in this or similar categories? Please list their website domains",
        "Who are your main competitors (list 3-5)? What is their strength, weakness and competitive edge? Please list their website domain",
        "What are the demographics/psychographics of your audience?",
        "Who are the audience that define/personify your brand the most? Who will readily buy the brand (early adopters)?",
        "Who are the audience that would never personify your brand?",
        "Who would represent/endorse your brand? Describe why?",
        "What are their pain/desires? What do they need out of the category?",
        "Who do you want to convince into buying your brand?",
        "How can we best reach them (main channels)?",
        "What is your brand purpose & values?",
        "What do you want people to say about your brand? What is the aspired brand personality? Which words do you want to own?",
        "Which traits will never define your brand?",
        "Is there any word, image, phrase or story that reflects what the brand stands for?",
        "Is there anything else you’d like to add?",
        "What are you expecting out of this branding exercise?",
        "List all your products/services/special features and the respective grouping if relevant.",
        "List any additions or extensions to your product/service that you might introduce in the future.",
    ]
    headers = [
        "The Brand Proposition",
        "The Category",
        "The Competition",
        "The Brand Purpose & Perception",
    ]

    def is_question(paragraph):
        return any(paragraph.strip() == q for q in questions)

    def is_header(paragraph):
        return any(paragraph.strip() == h for h in headers)

    for i in range(len(paragraphs)):
        paragraph_text = paragraphs[i].strip()
        next_paragraph_text = paragraphs[i + 1].strip() if i + 1 < len(paragraphs) else ""

        if is_question(paragraph_text) or (
                current_question and is_question(paragraph_text + " " + next_paragraph_text)):
            if current_question:
                data.append((current_question, '\n'.join(answer_lines)))
                answer_lines = []
            current_question = paragraph_text
            if next_paragraph_text and not is_question(next_paragraph_text) and not is_header(next_paragraph_text):
                current_question += " " + next_paragraph_text
        elif current_question and not is_header(paragraph_text):
            if paragraph_text not in current_question:
                answer_lines.append(paragraph_text)

    if current_question and answer_lines:
        data.append((current_question, '\n'.join(answer_lines)))

    data = [(q, a) for q, a in data if q and a]

    return data


st.title('Marketing Strategy and Campaign Generator')

openai_key = st.text_input("Enter your OpenAI API Key", type="password")
splitted_documents = []
if st.button('Confirm API Key'):
    if openai_key:
        st.session_state['openai_key'] = openai_key
        st.session_state['key_confirmed'] = True
        st.success('API Key confirmed. Please proceed.')
        os.environ["OPENAI_API_KEY"] = st.session_state['openai_key']
        openai.api_key = st.session_state['openai_key']
    else:
        st.error("Please enter the OpenAI key to proceed.")

if st.session_state.get('key_confirmed', False):
    uploaded_file = st.file_uploader("Choose a Word file", type=["docx"])
    BRAND_NAME = st.text_input("Enter your Brand Name")
    INDUSTRY = st.text_input("Enter your Industry")

    if st.button('Generate Marketing Strategy and Campaign'):
        if uploaded_file and BRAND_NAME and INDUSTRY:
            document_content = read_word_document(uploaded_file)

            for i, (question, answer) in enumerate(document_content):
                doc_id = f"doc_{i}"
                splitted_documents.append(
                    Document(page_content=question + "\n" + answer + "\n\n", metadata={"id": doc_id}))

            turbo = ChatOpenAI(model="gpt-3.5-turbo-16k", temperature=0.0, max_tokens=4000, streaming=False)

            embeddings = OpenAIEmbeddings(chunk_size=500)
            marketing_retriever = Chroma.from_documents(splitted_documents, embeddings, collection_name="brand")
            marketing_qa_template = """ the user is a marketing agent and he needs answers to some questions about your brand, the questions, and answers in the below context, so provide the answer to his question from the below context as it is in an accurate way, and don't make up answers or add info to the context answer from your mind.

            ===========================

            {context}

            Question: {question}"""

            MARKETING_PROMPT = PromptTemplate(
                template=marketing_qa_template, input_variables=["context", "question"]
            )
            marketing_qa = RetrievalQA.from_chain_type(
                llm=turbo,
                chain_type="stuff",
                retriever=marketing_retriever.as_retriever(search_kwargs={'k': 1}),
                chain_type_kwargs={"prompt": MARKETING_PROMPT, 'verbose': False},
            )

            questionaire_report = ""
            with st.spinner('Reading uploaded document...'):
                for i, q_a in enumerate(document_content):
                    questionaire_report += f"Q{i + 1}- " + q_a[0] + "\n" + marketing_qa.invoke(q_a[0])['result'] + "\n\n"

            marketing_strategy_template = f"""You are a brand called '{BRAND_NAME}', focusing on innovation and leadership within the '{INDUSTRY}' sector. Based on the provided answers to the questionnaire, generate a detailed marketing strategy and campaign.

            1. **Brand Insights**: Analyze the market dilemma, solutions, product offerings, and credibility. Identify core opportunities and how the brand differentiates itself.
            2. **Competitor Analysis**: Examine main competitors, their strengths, weaknesses, and market positions. Highlight the brand's competitive advantage.
            3. **Target Audience**: Define primary and secondary target audiences based on the demographics and psychographics. Outline strategies for engaging these audiences.
            4. **Marketing Strategy**:
                - **Brand Positioning**: Summarize the brand's market position and unique value proposition.
                - **Engagement Channels**: Identify the most effective channels for reaching the target audience.
                - **Content and Messaging**: Suggest key messages that resonate with the brand promise and audience's expectations.
            5. **Campaign Concept**: Propose a campaign that embodies strategic goals, including creative titles, objectives, key activities, and expected outcomes.
            6. **Implementation Plan**: Outline steps for executing the marketing strategy and campaign, including timelines and key milestones.
            7. **Evaluation Metrics**: Specify how to measure the success of the marketing strategy and campaign.

            Generate the marketing strategy and campaign in a cohesive narrative format."""

            main_model = ChatOpenAI(model="gpt-3.5-turbo", temperature=0.05, max_tokens=3000, streaming=False)
            prompt = PromptTemplate(template=marketing_strategy_template, input_variables=["context"])

            response = main_model({ "context": questionaire_report, "template": prompt.template })

            st.write(response)
        else:
            if not uploaded_file:
                st.error("Please upload one docx document.")
            if not BRAND_NAME:
                st.error("Please enter the brand name.")
            if not INDUSTRY:
                st.error("Please enter the industry of your brand.")
